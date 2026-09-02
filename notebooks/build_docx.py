"""Builds entregables/resumen_ejecutivo.docx from the English executive summary.

Same pattern as the earlier cases: the content is written HERE rather than parsed from the
Markdown, so the two stay in sync deliberately instead of by regex accident. Colours follow
this case's theme (crimson) rather than case 1's blue or case 2's green.

Note for whoever maintains case-template: this file is NOT in the template, even though the
template's requirements.txt already declares python-docx for it. It was copied across from
football-transfer-market and adapted. Fixing the template is its own job.

Run: python notebooks/build_docx.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

RUTA_BASE = Path(__file__).resolve().parents[1]
GRAFICOS = RUTA_BASE / 'salidas' / 'graficos'
OUT = RUTA_BASE / 'entregables' / 'resumen_ejecutivo.docx'

CARMESI = RGBColor(0xA8, 0x20, 0x3A)        # el acento del caso
CARMESI_OSCURO = RGBColor(0x6E, 0x15, 0x26)
DARKTEXT = RGBColor(0x1A, 0x18, 0x15)
GRAY = RGBColor(0x6E, 0x7B, 0x77)

doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].font.color.rgb = DARKTEXT


def heading(text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = CARMESI_OSCURO
        run.font.name = 'Cambria'
    return h


def para(text, *, italic=False, bold=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color or DARKTEXT
    return p


def bullet(texto_fuerte, resto):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(texto_fuerte)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = DARKTEXT
    r2 = p.add_run(resto)
    r2.font.size = Pt(11)
    r2.font.color.rgb = DARKTEXT
    return p


def image(filename, width_in=6.0, caption=None):
    doc.add_picture(str(GRAFICOS / filename), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = GRAY


# Title block ---------------------------------------------------------------
title = doc.add_paragraph()
r = title.add_run('Small-business credit: the vintage you are buying is younger than its number')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = CARMESI_OSCURO
r.font.name = 'Cambria'

sub = doc.add_paragraph()
r = sub.add_run('Executive summary for the investment committee  ·  2 September 2026')
r.font.size = Pt(12)
r.font.color.rgb = GRAY

doc.add_paragraph()

# What we set out to test ---------------------------------------------------
heading('What we set out to test')
para('The fund has to commit an expected-loss assumption for a book of US small-business loans '
     'now. The only evidence available for that decision is the charge-off rate of recent '
     'origination years — and we suspected that evidence is systematically flattering, because a '
     'loan approved eighteen months ago has not had time to default.')
para('We wrote three expectations down before opening the data. Two of them turned out to be '
     'wrong, and the way they were wrong is the useful part.', bold=True)

# Findings ------------------------------------------------------------------
heading('What we found')

para('A three-year-old vintage has shown you less than a quarter of what it will cost.', bold=True)
para('Across the twenty-four vintages old enough to have finished, only 23.5% of eventual '
     'charge-offs have occurred by month 36. The median default lands at month 58 — year five, not '
     'year three. Half the loss is still ahead at the five-year mark.')
image('02_a_los_tres_anos.png',
      caption='Share of eventual charge-offs already recorded, by months since approval.')

para('The 2023 vintage is on track to be the worst since 2010.', bold=True)
para('It has recorded 0.75 cents of loss per dollar approved, which reads as excellent. Corrected '
     'for age it projects to 4.24 cents, with an interquartile band of 3.28 to 6.18. Thirteen '
     'consecutive vintages, 2011 through 2022, landed below that. In loan-count terms the same '
     'vintage moves from 3.5% to 12.2%. The 2024 and 2025 books are published as a labelled gap '
     'rather than a number: they have realised 5.9% and 0.2% of their eventual loss, and '
     'multiplying those up would produce a confident-looking figure resting on almost nothing.')
image('01_lo_que_falta_por_llegar.png',
      caption='Loss already recorded and, above it, what the maturation of completed vintages '
              'implies is still to come.')

para('2020 and 2021 look like the best underwriting in a decade. They are not.', bold=True)
para('At the same age they default at roughly half the rate of every neighbouring vintage — 0.85% '
     'and 0.77% against 1.58% to 1.98%. Under CARES Act section 1112 the government was paying '
     'instalments on many of these loans. A model calibrated on those two years understates '
     'everything else by about 18%.')
image('04_artefacto_de_politica.png',
      caption='Identical axes: same age, half the default rate, for a reason that is not credit.')

para('What the correction does not do is reshuffle the league table.', bold=True)
para('The five worst vintages are the same whether you read raw rates or rates at equal age. The '
     'mistake a fund makes is not picking the wrong worst year — it is believing a young year’s '
     'number.')
image('03_nivel_no_orden.png',
      caption='Rank of each vintage under four readings. The order barely moves.')

para('The loss concentrates in short-term, small-ticket lending.', bold=True)
para('At 60 months, loans of seven years or less charge off at 5.67% against 0.39% for loans over '
     'twenty years. That is where the loss sits — though term is a proxy for product and '
     'collateral, so it is a statement about which business you are in, not a lever to pull.')

# Recommendations -----------------------------------------------------------
heading('What we recommend')
bullet('Price the book off a maturation curve, not off its observed rate. ',
       'For a 2023-like book the loss input moves from 0.75 to 4.24 cents per dollar approved. Low '
       'effort: a change of input plus a quarterly re-run.')
bullet('Keep 2020 and 2021 out of any calibration sample, ',
       'and use 2016–2019 as the modern reference window. Low effort.')
bullet('Underwrite the short-term, small-ticket segment on its own terms ',
       'instead of at a portfolio average. Medium effort: it touches origination policy, not a '
       'single number.')

# Limitations ---------------------------------------------------------------
heading('What this analysis cannot tell you')
para('These are SBA-guaranteed loans. The shape of the maturation curve transfers to unguaranteed '
     'private credit; the level does not, and nothing here can calibrate that gap. Every figure is '
     'gross loss — the source records no recoveries, so the fund’s actual loss is lower by an '
     'unknown amount. Nothing here is causal: the case describes how vintages behave, not what '
     'moves them. And the projection assumes the loss timing of 1991–2014 still holds; if today’s '
     'borrowers fail on a different schedule, the method is wrong in a way it cannot detect on its '
     'own.')
para('One check is open rather than passed. SBA publishes aggregate performance tables that would '
     'confirm our population in a single line, and the archive their own page links to returns a '
     '404. In its place the population was recounted independently from the raw files and checked '
     'against the programme’s published rules. That is internal consistency, not external '
     'reconciliation, and it stays on this list until it is closed.')
para('Full method, checks and decision log: CASO.md in the case repository.',
     italic=True, size=10, color=GRAY)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f'escrito {OUT.relative_to(RUTA_BASE)}  ({OUT.stat().st_size / 1024:.0f} KB)')
